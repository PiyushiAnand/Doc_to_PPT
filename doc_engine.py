from docx import Document

def generate_citation_doc(ppt_data, fact_registry, filename):
    """
    ppt_data: slide JSON used to generate PPT
    fact_registry: dict { fact_id -> {text, source} }
    """

    doc = Document()
    doc.add_heading("Citation Reference Document", level=0)

    used_fact_ids = set()

    # Collect all referenced fact IDs
    for slide in ppt_data.get("slides", []):
        for bullet in slide.get("bullets", []):
            for fid in bullet.get("fact_ids", []):
                used_fact_ids.add(fid)

    if not used_fact_ids:
        doc.add_paragraph("No citations found.")
        doc.save(filename)
        return

    # Write citations
    for fid in sorted(used_fact_ids):
        fact = fact_registry.get(fid, {})

        p = doc.add_paragraph()
        p.add_run(f"Fact ID: {fid}\n").bold = True
        p.add_run(f"Source: {fact.get('source', 'Unknown')}\n")
        p.add_run(f"Excerpt: \"{fact.get('text', '')}\"")
        doc.add_paragraph("-" * 30)

    doc.save(filename)
    print(f"Citation Doc Saved: {filename}")
